"""Traceable source-type extraction and processing provenance services."""

from __future__ import annotations

from dataclasses import asdict, dataclass
from datetime import datetime, timezone
import hashlib
from io import BytesIO
import json
from pathlib import Path
import sqlite3
import subprocess
import sys
from tempfile import TemporaryDirectory
from typing import Callable
from uuid import uuid4
import zipfile

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from metro_forensics.ids import stable_id
from metro_forensics.ingest import intake_evidence_root
from metro_forensics.review import open_review_task


LOW_CONFIDENCE_OCR_THRESHOLD = 80.0


@dataclass(frozen=True)
class ExtractedUnit:
    unit_kind: str
    unit_locator: str
    text: str
    confidence: float | None


@dataclass(frozen=True)
class ExtractionResult:
    units: list[ExtractedUnit]
    warnings: list[str]

    @property
    def text(self) -> str:
        return "\n".join(unit.text for unit in self.units if unit.text)

    def source_location_mapping(self) -> str:
        return json.dumps([asdict(unit) for unit in self.units], ensure_ascii=False, sort_keys=True)


@dataclass(frozen=True)
class ProcessingResult:
    processing_run_id: str
    derivative_id: str


@dataclass(frozen=True)
class UnsupportedResult:
    processing_run_id: str
    review_task_id: str


def extract_pdf(data: bytes) -> ExtractionResult:
    """Extract each PDF page without losing its one-based page locator."""
    reader = PdfReader(BytesIO(data))
    units: list[ExtractedUnit] = []
    warnings: list[str] = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        confidence: float | None = None
        if not text:
            text, confidence, ocr_warning = _ocr_pdf_page(data, page_number)
            if ocr_warning:
                warnings.append(ocr_warning)
            if confidence is None or confidence < LOW_CONFIDENCE_OCR_THRESHOLD:
                display_confidence = "unavailable" if confidence is None else f"{confidence:.2f}"
                warnings.append(
                    f"LOW_CONFIDENCE_OCR page:{page_number} confidence:{display_confidence}"
                )
        units.append(ExtractedUnit("PAGE", f"page:{page_number}", text, confidence))
    return ExtractionResult(units, warnings)


def extract_xlsx(data: bytes) -> ExtractionResult:
    """Extract non-empty cells while retaining formulas, not cached display values."""
    workbook = load_workbook(BytesIO(data), read_only=True, data_only=False)
    units: list[ExtractedUnit] = []
    try:
        for worksheet in workbook.worksheets:
            for row in worksheet.iter_rows():
                for cell in row:
                    if cell.value is None:
                        continue
                    units.append(ExtractedUnit(
                        "CELL", f"sheet:{worksheet.title}!{cell.coordinate}", str(cell.value), None
                    ))
    finally:
        workbook.close()
    return ExtractionResult(units, [])


def extract_docx(data: bytes) -> ExtractionResult:
    """Extract DOCX paragraphs and table cells with independent locators."""
    document = Document(BytesIO(data))
    units: list[ExtractedUnit] = []
    for paragraph_number, paragraph in enumerate(document.paragraphs, start=1):
        if paragraph.text:
            units.append(ExtractedUnit(
                "PARAGRAPH", f"paragraph:{paragraph_number}", paragraph.text, None
            ))
    for table_number, table in enumerate(document.tables, start=1):
        for row_number, row in enumerate(table.rows, start=1):
            for cell_number, cell in enumerate(row.cells, start=1):
                if cell.text:
                    units.append(ExtractedUnit(
                        "TABLE_CELL", f"table:{table_number},row:{row_number},cell:{cell_number}",
                        cell.text, None,
                    ))
    return ExtractionResult(units, [])


def record_processing_result(
    conn: sqlite3.Connection, source_file_id: str, operation: str, derivative_bytes: bytes
) -> ProcessingResult:
    """Record an already-created derivative's provenance without touching evidence files."""
    with conn:
        processing_run_id = _start_processing_run(conn, source_file_id, operation, {})
        _finish_processing_run(conn, processing_run_id)
        derivative_id = _insert_derivative(
            conn,
            source_file_id,
            processing_run_id,
            derivative_bytes,
            "RECORDED_BYTES",
            f"memory://{source_file_id}/{processing_run_id}",
            "[]",
        )
    return ProcessingResult(processing_run_id, derivative_id)


def record_unsupported_legacy_doc(conn: sqlite3.Connection, source_file_id: str) -> UnsupportedResult:
    """Preserve an unextractable legacy DOC as an open human-review exception."""
    with conn:
        source = _source_row(conn, source_file_id)
        processing_run_id = _start_processing_run(
            conn, source_file_id, "EXTRACT_DOC", {"reason": "UNSUPPORTED_LEGACY_DOC"}
        )
        _finish_processing_run(conn, processing_run_id, errors="UNSUPPORTED_LEGACY_DOC")
        review_task_id = _open_review_task(
            conn,
            source["package_id"],
            source_file_id,
            "EXTRACTION_EXCEPTION",
            "UNSUPPORTED_LEGACY_DOC",
            "Legacy .doc was preserved without pseudo-text extraction; human review is required.",
        )
    return UnsupportedResult(processing_run_id, review_task_id)


def process_source(
    conn: sqlite3.Connection, source_file_id: str, evidence_root: Path, derivative_root: Path
) -> list[str]:
    """Extract one immutable archive member into a new source-mapped derivative."""
    evidence_path = intake_evidence_root(conn).resolve()
    supplied_evidence_path = evidence_root.resolve()
    if supplied_evidence_path != evidence_path:
        raise ValueError("evidence root does not match the ledger's immutable intake root")
    derivative_path = derivative_root.resolve()
    if (
        evidence_path == derivative_path
        or evidence_path in derivative_path.parents
        or derivative_path in evidence_path.parents
    ):
        raise ValueError("derivative root must not overlap immutable evidence root")
    source = _source_row(conn, source_file_id)
    suffix = Path(source["archive_member_path"]).suffix.lower()
    if suffix == ".doc":
        record_unsupported_legacy_doc(conn, source_file_id)
        return []

    adapters: dict[str, tuple[str, Callable[[bytes], ExtractionResult]]] = {
        ".pdf": ("EXTRACT_PDF", extract_pdf),
        ".xlsx": ("EXTRACT_XLSX", extract_xlsx),
        ".xlsm": ("EXTRACT_XLSX", extract_xlsx),
        ".docx": ("EXTRACT_DOCX", extract_docx),
    }
    if suffix not in adapters:
        return _record_failed_processing(
            conn, source, "EXTRACT_UNSUPPORTED", "UNSUPPORTED_SOURCE_TYPE"
        )

    operation, adapter = adapters[suffix]
    with conn:
        processing_run_id = _start_processing_run(
            conn, source_file_id, operation, {"archive_member_path": source["archive_member_path"]}
        )
    artifact_path: Path | None = None
    try:
        data = _read_source_bytes(source, evidence_path)
        actual_hash = hashlib.sha256(data).hexdigest()
        if actual_hash != source["sha256"]:
            raise ValueError("SOURCE_HASH_MISMATCH")
        extracted = adapter(data)
        artifact_bytes = _serialized_derivative(source_file_id, processing_run_id, operation, extracted)
        artifact_path = _write_derivative(
            derivative_root, source_file_id, processing_run_id, artifact_bytes
        )
        with conn:
            _finish_processing_run(conn, processing_run_id, warnings="\n".join(extracted.warnings))
            derivative_id = _insert_derivative(
                conn,
                source_file_id,
                processing_run_id,
                artifact_bytes,
                "EXTRACTED_TEXT_JSON",
                str(artifact_path),
                extracted.source_location_mapping(),
            )
            low_confidence_warnings = [
                warning for warning in extracted.warnings if warning.startswith("LOW_CONFIDENCE_OCR")
            ]
            if low_confidence_warnings:
                _open_review_task(
                    conn,
                    source["package_id"],
                    source_file_id,
                    "EXTRACTION_EXCEPTION",
                    "LOW_CONFIDENCE_OCR",
                    "OCR output requires human verification: " + "; ".join(low_confidence_warnings),
                )
        return [derivative_id]
    except Exception as exc:
        error = str(exc) or type(exc).__name__
        if artifact_path is not None:
            try:
                _remove_derivative_artifact(artifact_path, derivative_root)
            except OSError as cleanup_error:
                error = f"{error}; DERIVATIVE_CLEANUP_FAILED: {cleanup_error}"
        with conn:
            _finish_processing_run(conn, processing_run_id, errors=error)
            _open_review_task(
                conn,
                source["package_id"],
                source_file_id,
                "EXTRACTION_EXCEPTION",
                "EXTRACTION_FAILED",
                f"{operation} failed: {error}",
            )
        return []


def _record_failed_processing(
    conn: sqlite3.Connection, source: sqlite3.Row, operation: str, reason_code: str
) -> list[str]:
    with conn:
        processing_run_id = _start_processing_run(conn, source["source_file_id"], operation, {})
        _finish_processing_run(conn, processing_run_id, errors=reason_code)
        _open_review_task(
            conn, source["package_id"], source["source_file_id"], "EXTRACTION_EXCEPTION",
            reason_code, f"No safe adapter is configured for {source['archive_member_path']}.",
        )
    return []


def _source_row(conn: sqlite3.Connection, source_file_id: str) -> sqlite3.Row:
    source = conn.execute(
        """SELECT source_file.source_file_id, source_file.package_id, source_file.archive_member_path,
                  source_file.sha256, package.production_archive_path
           FROM source_file JOIN package USING (package_id)
           WHERE source_file_id=?""",
        (source_file_id,),
    ).fetchone()
    if source is None:
        raise ValueError(f"unknown source_file_id: {source_file_id}")
    return source


def _start_processing_run(
    conn: sqlite3.Connection, source_file_id: str, operation: str, parameters: dict[str, object]
) -> str:
    _source_row(conn, source_file_id)
    nonce = uuid4().hex
    started_at = _now()
    processing_run_id = stable_id("PR", source_file_id, operation, started_at, nonce)
    conn.execute(
        """INSERT INTO processing_run(
               processing_run_id, source_file_id, operation, tool_name, tool_version, parameters, started_at
           ) VALUES(?,?,?,?,?,?,?)""",
        (
            processing_run_id, source_file_id, operation, "metro_forensics.extract",
            sys.version.split()[0], json.dumps(parameters, sort_keys=True), started_at,
        ),
    )
    return processing_run_id


def _finish_processing_run(
    conn: sqlite3.Connection, processing_run_id: str, warnings: str = "", errors: str = ""
) -> None:
    conn.execute(
        "UPDATE processing_run SET completed_at=?, warnings=?, errors=? WHERE processing_run_id=?",
        (_now(), warnings, errors, processing_run_id),
    )


def _insert_derivative(
    conn: sqlite3.Connection,
    source_file_id: str,
    processing_run_id: str,
    derivative_bytes: bytes,
    artifact_type: str,
    artifact_path: str,
    source_location_mapping: str,
) -> str:
    digest = hashlib.sha256(derivative_bytes).hexdigest()
    derivative_id = stable_id("DV", source_file_id, processing_run_id, digest)
    conn.execute(
        """INSERT INTO derivative(
               derivative_id, source_file_id, processing_run_id, sha256, artifact_type, artifact_path,
               source_location_mapping
           ) VALUES(?,?,?,?,?,?,?)""",
        (
            derivative_id, source_file_id, processing_run_id, digest, artifact_type, artifact_path,
            source_location_mapping,
        ),
    )
    return derivative_id


def _open_review_task(
    conn: sqlite3.Connection,
    package_id: str,
    source_file_id: str,
    task_type: str,
    reason_code: str,
    concern: str,
) -> str:
    return open_review_task(
        conn,
        "SOURCE_FILE",
        source_file_id,
        reason_code,
        concern,
        task_type=task_type,
        task_key=uuid4().hex,
    )


def _read_source_bytes(source: sqlite3.Row, evidence_root: Path) -> bytes:
    archive_name = source["production_archive_path"]
    if not archive_name:
        raise ValueError("MISSING_PRODUCTION_ARCHIVE")
    root = evidence_root.resolve()
    archive_path = (root / archive_name).resolve()
    if not archive_path.is_relative_to(root):
        raise ValueError("UNSAFE_PRODUCTION_ARCHIVE_PATH")
    with zipfile.ZipFile(archive_path) as archive:
        return archive.read(source["archive_member_path"])


def _write_derivative(
    derivative_root: Path, source_file_id: str, processing_run_id: str, payload: bytes
) -> Path:
    root = derivative_root.resolve()
    if not _safe_path_component(source_file_id) or not _safe_path_component(processing_run_id):
        raise ValueError("UNSAFE_DERIVATIVE_PATH")
    artifact_path = (root / source_file_id / processing_run_id / "extracted.json").resolve()
    if not artifact_path.is_relative_to(root):
        raise ValueError("UNSAFE_DERIVATIVE_PATH")
    artifact_path.parent.mkdir(parents=True, exist_ok=False)
    artifact_path.write_bytes(payload)
    return artifact_path


def _safe_path_component(value: str) -> bool:
    return value not in {"", ".", ".."} and Path(value).name == value


def _remove_derivative_artifact(artifact_path: Path, derivative_root: Path) -> None:
    """Remove a failed write only after confirming it belongs to this derivative root."""
    root = derivative_root.resolve()
    candidate = artifact_path.resolve()
    if not candidate.is_relative_to(root):
        raise OSError("refusing to clean a path outside derivative_root")
    candidate.unlink(missing_ok=True)
    for directory in (candidate.parent, candidate.parent.parent):
        try:
            directory.rmdir()
        except OSError:
            break


def _serialized_derivative(
    source_file_id: str, processing_run_id: str, operation: str, extracted: ExtractionResult
) -> bytes:
    return json.dumps(
        {
            "source_file_id": source_file_id,
            "processing_run_id": processing_run_id,
            "operation": operation,
            "text": extracted.text,
            "units": [asdict(unit) for unit in extracted.units],
            "warnings": extracted.warnings,
        },
        ensure_ascii=False,
        sort_keys=True,
    ).encode("utf-8")


def _ocr_pdf_page(data: bytes, page_number: int) -> tuple[str, float | None, str | None]:
    """OCR one page only after embedded-text extraction was unusable."""
    with TemporaryDirectory() as temporary_directory:
        temporary_path = Path(temporary_directory)
        input_path = temporary_path / "source.pdf"
        page_prefix = temporary_path / "page"
        input_path.write_bytes(data)
        try:
            subprocess.run(
                ["pdftoppm", "-f", str(page_number), "-l", str(page_number), "-png",
                 str(input_path), str(page_prefix)],
                check=True, capture_output=True, text=True,
            )
            image_path = temporary_path / f"page-{page_number}.png"
            recognized = subprocess.run(
                ["tesseract", str(image_path), "stdout", "tsv"],
                check=True, capture_output=True, text=True,
            )
        except (FileNotFoundError, subprocess.CalledProcessError) as exc:
            return "", None, f"OCR_UNAVAILABLE page:{page_number}: {exc}"
    rows = [line.split("\t") for line in recognized.stdout.splitlines()[1:] if line]
    words = [row[11] for row in rows if len(row) > 11 and row[11].strip()]
    confidences = [float(row[10]) for row in rows if len(row) > 11 and row[11].strip() and row[10] != "-1"]
    confidence = sum(confidences) / len(confidences) if confidences else None
    return " ".join(words), confidence, None


def _now() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="microseconds").replace("+00:00", "Z")
