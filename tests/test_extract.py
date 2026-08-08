import hashlib
import io
import json
import sqlite3
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen.canvas import Canvas

from metro_forensics.extract import (
    extract_docx,
    extract_pdf,
    extract_xlsx,
    process_source,
    record_processing_result,
    record_unsupported_legacy_doc,
)
from metro_forensics.ingest import INTAKE_EVIDENCE_ROOT_KEY
from tests.helpers import new_test_db, seed_package_source


class ExtractionTests(unittest.TestCase):
    def test_derivative_requires_a_successfully_completed_processing_run(self):
        """Fail if a failed or unfinished attempt can receive a derivative."""
        db = new_test_db()
        _, source_id = seed_package_source(db)
        db.execute(
            "INSERT INTO processing_run(processing_run_id,source_file_id,operation,tool_name,started_at) "
            "VALUES('PR_UNFINISHED',?,'EXTRACT_XLSX','test','2026-08-07T12:00:00Z')",
            (source_id,),
        )

        with self.assertRaises(sqlite3.IntegrityError):
            db.execute(
                "INSERT INTO derivative(derivative_id,source_file_id,processing_run_id,sha256,artifact_type,artifact_path,source_location_mapping) "
                "VALUES('DV_UNFINISHED',?,'PR_UNFINISHED',?,'RECORDED_BYTES','memory://unfinished','[]')",
                (source_id, "1" * 64),
            )

    def test_source_provenance_is_immutable_after_processing_exists(self):
        """Fail if an existing run can be relabeled with new source bytes or coordinates."""
        db = new_test_db()
        _, source_id = seed_package_source(db)
        record_processing_result(db, source_id, "EXTRACT_XLSX", b"first")

        for field, value in (("sha256", "1" * 64), ("archive_member_path", "rewritten.xlsx"), ("byte_size", 2)):
            with self.subTest(field=field), self.assertRaises(sqlite3.IntegrityError):
                db.execute(f"UPDATE source_file SET {field}=? WHERE source_file_id=?", (value, source_id))

    def test_persistence_failure_removes_unledgered_derivative_artifact(self):
        """Fail if a database failure leaves an extracted file without a DERIVATIVE row."""
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            evidence_root = root / "evidence"
            derivative_root = root / "analysis" / "derivatives"
            evidence_root.mkdir()
            source_bytes = self._workbook_bytes_with_formula()
            with zipfile.ZipFile(evidence_root / "production.zip", "w") as archive:
                archive.writestr("financials.xlsx", source_bytes)
            db = new_test_db()
            _, source_id = seed_package_source(db, member="financials.xlsx")
            db.execute(
                "UPDATE package SET production_archive_path='production.zip' WHERE package_id='P1'"
            )
            db.execute(
                "UPDATE source_file SET sha256=?, byte_size=? WHERE source_file_id=?",
                (hashlib.sha256(source_bytes).hexdigest(), len(source_bytes), source_id),
            )
            db.execute(
                "CREATE TRIGGER reject_derivative_for_test BEFORE INSERT ON derivative "
                "BEGIN SELECT RAISE(ABORT, 'forced derivative failure'); END"
            )
            self._bind_evidence_root(db, evidence_root)

            derivative_ids = process_source(db, source_id, evidence_root, derivative_root)

            self.assertEqual([], derivative_ids)
            self.assertEqual([], list(derivative_root.rglob("extracted.json")))
            self.assertEqual(0, db.execute("SELECT count(*) FROM derivative").fetchone()[0])

    def test_reprocessing_creates_new_run_and_derivative_ids(self):
        """Fail if a retry overwrites its earlier processing provenance."""
        db = new_test_db()
        _, source_id = seed_package_source(db, member="sheet.xlsx")

        first = record_processing_result(db, source_id, "EXTRACT_XLSX", b"first")
        second = record_processing_result(db, source_id, "EXTRACT_XLSX", b"first")

        self.assertNotEqual(first.processing_run_id, second.processing_run_id)
        self.assertNotEqual(first.derivative_id, second.derivative_id)
        self.assertEqual(2, db.execute(
            "SELECT count(*) FROM processing_run WHERE source_file_id=?", (source_id,)
        ).fetchone()[0])
        self.assertEqual(2, db.execute(
            "SELECT count(*) FROM derivative WHERE source_file_id=?", (source_id,)
        ).fetchone()[0])
        self.assertEqual("0" * 64, db.execute(
            "SELECT sha256 FROM source_file WHERE source_file_id=?", (source_id,)
        ).fetchone()[0])

    def test_legacy_doc_records_attempt_and_opens_review_exception(self):
        """Fail if an unsupported DOC is silently converted into text."""
        db = new_test_db()
        _, source_id = seed_package_source(db, member="legacy.doc")

        result = record_unsupported_legacy_doc(db, source_id)

        task = db.execute(
            "SELECT task_state,reason_code FROM review_task WHERE review_task_id=?",
            (result.review_task_id,),
        ).fetchone()
        run = db.execute(
            "SELECT operation,errors FROM processing_run WHERE processing_run_id=?",
            (result.processing_run_id,),
        ).fetchone()
        self.assertEqual(("OPEN", "UNSUPPORTED_LEGACY_DOC"), tuple(task))
        self.assertEqual("EXTRACT_DOC", run["operation"])
        self.assertIn("UNSUPPORTED_LEGACY_DOC", run["errors"])
        self.assertEqual(0, db.execute(
            "SELECT count(*) FROM derivative WHERE source_file_id=?", (source_id,)
        ).fetchone()[0])

    def test_derivative_cannot_claim_a_different_source_than_its_run(self):
        """Fail if a derivative can forge a source/run provenance chain."""
        db = new_test_db()
        _, first_source_id = seed_package_source(db, source_file_id="S1", member="first.xlsx")
        db.execute(
            "INSERT INTO source_file(source_file_id,package_id,archive_member_path,byte_size,sha256,media_type) "
            "VALUES('S2','P1','second.xlsx',1,?,'application/octet-stream')",
            ("1" * 64,),
        )
        first = record_processing_result(db, first_source_id, "EXTRACT_XLSX", b"first")

        with self.assertRaises(sqlite3.IntegrityError):
            db.execute(
                "INSERT INTO derivative(derivative_id,source_file_id,processing_run_id,sha256,artifact_type,artifact_path,source_location_mapping) "
                "VALUES('DV_FORGED','S2',?,?, 'RECORDED_BYTES','memory://forged','[]')",
                (first.processing_run_id, "2" * 64),
            )

    def test_completed_processing_provenance_cannot_be_rewritten(self):
        """Fail if a completed run can be moved after it has a derivative."""
        db = new_test_db()
        _, first_source_id = seed_package_source(db, source_file_id="S1", member="first.xlsx")
        db.execute(
            "INSERT INTO source_file(source_file_id,package_id,archive_member_path,byte_size,sha256,media_type) "
            "VALUES('S2','P1','second.xlsx',1,?,'application/octet-stream')",
            ("1" * 64,),
        )
        result = record_processing_result(db, first_source_id, "EXTRACT_XLSX", b"first")

        with self.assertRaises(sqlite3.IntegrityError):
            db.execute(
                "UPDATE processing_run SET source_file_id='S2' WHERE processing_run_id=?",
                (result.processing_run_id,),
            )
        with self.assertRaises(sqlite3.IntegrityError):
            db.execute("DELETE FROM derivative WHERE derivative_id=?", (result.derivative_id,))

    def test_process_source_rejects_traversal_in_provenance_id_before_writing(self):
        """Fail if a forged source ID can make a derivative escape its supplied root."""
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            evidence_root = root / "evidence"
            derivative_root = root / "analysis" / "derivatives"
            evidence_root.mkdir()
            source_bytes = self._workbook_bytes_with_formula()
            with zipfile.ZipFile(evidence_root / "production.zip", "w") as archive:
                archive.writestr("financials.xlsx", source_bytes)
            db = new_test_db()
            db.execute(
                "INSERT INTO package(package_id,control_record_path,production_archive_path,expected_level1_count) "
                "VALUES('P1','P1.pdf','production.zip',1)"
            )
            db.execute(
                "INSERT INTO source_file(source_file_id,package_id,archive_member_path,byte_size,sha256,media_type) "
                "VALUES(?,?,?,?,?,?)",
                (
                    "safe/../../escaped", "P1", "financials.xlsx", len(source_bytes),
                    hashlib.sha256(source_bytes).hexdigest(),
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                ),
            )
            self._bind_evidence_root(db, evidence_root)

            result = process_source(db, "safe/../../escaped", evidence_root, derivative_root)

            self.assertEqual([], result)
            self.assertFalse((root / "escaped").exists())
            self.assertEqual(0, db.execute("SELECT count(*) FROM derivative").fetchone()[0])

    def test_blank_pdf_ocr_opens_low_confidence_review_task(self):
        """Fail if OCR with no usable confidence is treated as verified-quality text."""
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            evidence_root = root / "evidence"
            derivative_root = root / "analysis" / "derivatives"
            evidence_root.mkdir()
            source_bytes = self._pdf_bytes("")
            with zipfile.ZipFile(evidence_root / "production.zip", "w") as archive:
                archive.writestr("scan.pdf", source_bytes)
            db = new_test_db()
            _, source_id = seed_package_source(db, member="scan.pdf")
            db.execute(
                "UPDATE package SET production_archive_path=? WHERE package_id='P1'", ("production.zip",)
            )
            db.execute(
                "UPDATE source_file SET sha256=?, byte_size=? WHERE source_file_id=?",
                (hashlib.sha256(source_bytes).hexdigest(), len(source_bytes), source_id),
            )
            self._bind_evidence_root(db, evidence_root)

            derivative_ids = process_source(db, source_id, evidence_root, derivative_root)

            self.assertEqual(1, len(derivative_ids))
            task = db.execute(
                "SELECT reason_code FROM review_task WHERE source_file_id=?", (source_id,)
            ).fetchone()
            self.assertIsNotNone(task)
            self.assertEqual("LOW_CONFIDENCE_OCR", task[0])

    def test_pdf_adapter_maps_embedded_text_to_its_page(self):
        """Fail if embedded PDF text loses its original page coordinate."""
        data = self._pdf_bytes("Metro response")

        result = extract_pdf(data)

        self.assertEqual("PAGE", result.units[0].unit_kind)
        self.assertEqual("page:1", result.units[0].unit_locator)
        self.assertIn("Metro response", result.units[0].text)
        self.assertEqual([], result.warnings)

    def test_xlsx_adapter_preserves_formula_as_formula_at_exact_cell(self):
        """Fail if a formula is replaced by a displayed/cached value."""
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Data"
        sheet["A1"] = "amount"
        sheet["B1"] = "=1+2"
        data = self._workbook_bytes(workbook)

        result = extract_xlsx(data)
        by_locator = {unit.unit_locator: unit for unit in result.units}

        self.assertEqual("amount", by_locator["sheet:Data!A1"].text)
        self.assertEqual("=1+2", by_locator["sheet:Data!B1"].text)

    def test_docx_adapter_maps_paragraphs_and_table_cells(self):
        """Fail if DOCX tables or paragraphs become untraceable flat text."""
        document = Document()
        document.add_paragraph("Opening paragraph")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Contract"
        table.cell(0, 1).text = "2026"
        stream = io.BytesIO()
        document.save(stream)

        result = extract_docx(stream.getvalue())
        by_locator = {unit.unit_locator: unit for unit in result.units}

        self.assertEqual("Opening paragraph", by_locator["paragraph:1"].text)
        self.assertEqual("Contract", by_locator["table:1,row:1,cell:1"].text)
        self.assertEqual("2026", by_locator["table:1,row:1,cell:2"].text)

    def test_process_source_writes_source_mapped_derivative_under_supplied_root(self):
        """Fail if processing writes outside its derivative root or omits mapping."""
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            evidence_root = root / "evidence"
            derivative_root = root / "analysis" / "derivatives"
            evidence_root.mkdir()
            source_bytes = self._workbook_bytes_with_formula()
            archive_path = evidence_root / "production.zip"
            with zipfile.ZipFile(archive_path, "w") as archive:
                archive.writestr("financials.xlsx", source_bytes)

            db = new_test_db()
            _, source_id = seed_package_source(db, member="financials.xlsx")
            db.execute(
                "UPDATE package SET production_archive_path=? WHERE package_id='P1'",
                ("production.zip",),
            )
            db.execute(
                "UPDATE source_file SET sha256=?, byte_size=?, media_type=? WHERE source_file_id=?",
                (
                    hashlib.sha256(source_bytes).hexdigest(), len(source_bytes),
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", source_id,
                ),
            )
            self._bind_evidence_root(db, evidence_root)

            derivative_ids = process_source(db, source_id, evidence_root, derivative_root)

            self.assertEqual(1, len(derivative_ids))
            derivative = db.execute(
                "SELECT processing_run_id,artifact_path,source_location_mapping FROM derivative "
                "WHERE derivative_id=?", (derivative_ids[0],)
            ).fetchone()
            artifact_path = Path(derivative["artifact_path"])
            self.assertTrue(artifact_path.is_relative_to(derivative_root.resolve()))
            self.assertTrue(artifact_path.is_file())
            mapping = json.loads(derivative["source_location_mapping"])
            self.assertEqual("sheet:Data!B1", mapping[1]["unit_locator"])
            self.assertEqual(1, db.execute(
                "SELECT count(*) FROM processing_run WHERE source_file_id=? AND processing_run_id=?",
                (source_id, derivative["processing_run_id"]),
            ).fetchone()[0])

    @staticmethod
    def _pdf_bytes(text):
        stream = io.BytesIO()
        canvas = Canvas(stream)
        canvas.drawString(72, 720, text)
        canvas.save()
        return stream.getvalue()

    @staticmethod
    def _workbook_bytes(workbook):
        stream = io.BytesIO()
        workbook.save(stream)
        return stream.getvalue()

    def _workbook_bytes_with_formula(self):
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Data"
        sheet["A1"] = "amount"
        sheet["B1"] = "=1+2"
        return self._workbook_bytes(workbook)

    @staticmethod
    def _bind_evidence_root(db, evidence_root):
        db.execute(
            "INSERT INTO operational_metadata(key,value) VALUES(?,?)",
            (INTAKE_EVIDENCE_ROOT_KEY, str(evidence_root.resolve())),
        )
